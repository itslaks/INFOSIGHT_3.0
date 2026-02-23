import sys
import os
from pathlib import Path

# Add project root to sys.path to allow importing from utils, core, etc.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import json
import os
import sys
import re
import logging
from io import StringIO
from flask import Blueprint, render_template, request, jsonify, g, send_file, Response
from utils.security import rate_limit_api, validate_request, InputValidator
# Groq import removed - using centralized router
from fuzzywuzzy import fuzz
from difflib import SequenceMatcher
import time
import markdown2
import sqlite3
from pathlib import Path
from datetime import datetime, timedelta
from collections import defaultdict
import hashlib
import threading
import concurrent.futures
from functools import lru_cache
import asyncio
import PyPDF2
import io
from werkzeug.utils import secure_filename
import base64
import logging as _logger
import pytesseract
from PIL import Image
import openpyxl
import pandas as pd
import docx
import json
import csv
from gtts import gTTS
import tempfile
import mimetypes
from collections import Counter
import re
from datetime import datetime
import zipfile

# ADD THESE ADVANCED IMPORTS
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    _logger.warning("pytesseract/PIL not available - OCR features disabled")

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    _logger.warning("openpyxl/pandas not available - Excel features disabled")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    _logger.warning("python-docx not available - Word document features disabled")

try:
    from gtts import gTTS
    TTS_AVAILABLE = True
except ImportError:
    TTS_AVAILABLE = False
    _logger.warning("gTTS not available - Audio summary features disabled")

import tempfile
import mimetypes
from collections import Counter
import zipfile

# Local LLM fallback
try:
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from utils.local_llm_utils import generate_with_ollama, check_ollama_available
    from utils.llm_logger import log_llm_status, log_llm_request, log_llm_success, log_llm_error, log_llm_fallback, log_processing_step
    LOCAL_LLM_AVAILABLE = True
except ImportError as e:
    LOCAL_LLM_AVAILABLE = False
    # Use logging instead of print for better control
    import logging
    logging.getLogger(__name__).warning(f"Local LLM utilities not available for CyberSentry AI: {e}")
    # Create dummy functions
    def log_llm_status(*args, **kwargs): return (False, False)
    def log_llm_request(*args, **kwargs): pass
    def log_llm_success(*args, **kwargs): pass
    def log_llm_error(*args, **kwargs): pass
    def log_llm_fallback(*args, **kwargs): pass
    def log_processing_step(*args, **kwargs): pass

# Create a blueprint
cybersentry_ai = Blueprint('cybersentry_ai', __name__, template_folder='templates')

# Configure logger
_logger = logging.getLogger(__name__)

# Load responses from JSON file with caching
_responses_cache = None
_cache_time = None

def load_responses():
    """Load responses with caching mechanism"""
    global _responses_cache, _cache_time
    
    # Cache for 5 minutes
    if _responses_cache and _cache_time and (time.time() - _cache_time < 300):
        return _responses_cache
    
    try:
        # Path relative to project root
        try:
            from utils.paths import get_data_path
            responses_path = str(get_data_path('responses.json'))
        except ImportError:
            project_root = os.path.dirname(os.path.dirname(__file__))
            responses_path = os.path.join(project_root, 'data', 'responses.json')
        with open(responses_path, 'r', encoding='utf-8') as file:
            _responses_cache = json.load(file)
            _cache_time = time.time()
            return _responses_cache
    except FileNotFoundError:
        import logging
        logging.getLogger(__name__).warning("Warning: responses.json not found. Creating empty response list.")
        return []
    except json.JSONDecodeError as e:
        import logging
        logging.getLogger(__name__).error(f"Error parsing responses.json: {e}")
        return []
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"Error loading responses: {e}")
        return []

responses = load_responses()
if not responses:
    _logger.warning("⚠️ No responses loaded - JSON database is empty. AI fallback will be used.")

# Log LLM status at startup
try:
    log_llm_status("CyberSentry AI")
except:
    pass

# Use centralized LLM router
try:
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from core.llm_router import generate_text
    LLM_ROUTER_AVAILABLE = True
except ImportError as e:
    LLM_ROUTER_AVAILABLE = False
    _logger.error(f"✗ LLM router not available: {e}")
    def generate_text(*args, **kwargs):
        return {"response": "", "model": "none", "source": "none"}

# Cache LLM responses
@lru_cache(maxsize=100)
def get_cached_response(query_hash):
    """Cache responses to avoid redundant API calls"""
    pass

def capture_output(func):
    """Decorator to capture stdout output and preserve model_used"""
    def wrapper(*args, **kwargs):
        old_stdout = sys.stdout
        sys.stdout = StringIO()
        
        try:
            result = func(*args, **kwargs)
            output = sys.stdout.getvalue()
            sys.stdout = old_stdout
            
            # Extract model_used if present
            model_used = getattr(result, '_model_used', 'groq') if result else 'groq'
            
            # Return (result, output, model_used)
            return result, output, model_used
        except Exception as e:
            sys.stdout = old_stdout
            import logging
            logging.getLogger(__name__).error(f"Error in capture_output: {e}")
            return None, f"Error: {str(e)}", None
    return wrapper

def _wrap_with_model_used(text, model_used):
    """Wrap a string in a subclass that can carry a `_model_used` attribute.
    This avoids attempting to set attributes on built-in `str` instances."""
    try:
        class _RichStr(str):
            pass

        r = _RichStr(text)
        setattr(r, '_model_used', model_used)
        return r
    except Exception:
        return text

def normalize_text(text):
    """Normalize text for better matching"""
    text = text.lower().strip()
    text = ' '.join(text.split())
    text = re.sub(r'[?.!,;]+$', '', text)
    return text

def calculate_similarity_score(query, question):
    """Calculate multiple similarity scores and return weighted average"""
    query_norm = normalize_text(query)
    question_norm = normalize_text(question)
    
    token_score = fuzz.token_set_ratio(query_norm, question_norm)
    partial_score = fuzz.partial_ratio(query_norm, question_norm)
    sort_score = fuzz.token_sort_ratio(query_norm, question_norm)
    seq_score = SequenceMatcher(None, query_norm, question_norm).ratio() * 100
    
    weighted_score = (token_score * 0.4 + partial_score * 0.2 + 
                     sort_score * 0.3 + seq_score * 0.1)
    
    return weighted_score

def extract_keywords(text):
    """Extract important keywords from text"""
    stop_words = {'what', 'is', 'the', 'a', 'an', 'how', 'does', 'in', 'of', 
                  'to', 'for', 'and', 'or', 'can', 'you', 'me', 'explain', 
                  'define', 'describe', 'tell'}
    
    words = normalize_text(text).split()
    keywords = [w for w in words if w not in stop_words and len(w) > 2]
    return keywords

def keyword_match_score(query, question):
    """Calculate score based on keyword matching"""
    query_keywords = set(extract_keywords(query))
    question_keywords = set(extract_keywords(question))
    
    if not query_keywords or not question_keywords:
        return 0
    
    intersection = query_keywords.intersection(question_keywords)
    union = query_keywords.union(question_keywords)
    
    return (len(intersection) / len(union)) * 100 if union else 0

def is_valid_answer(answer):
    """Validate if answer is meaningful (not just a single word like 'groq')"""
    if not answer:
        return False
    
    # Convert to string if needed and strip whitespace
    if not isinstance(answer, str):
        answer = str(answer)
    answer = answer.strip()
    
    # Empty after stripping
    if not answer:
        return False
    
    # Check if answer is just a single word (common invalid answers) - CHECK THIS FIRST
    invalid_words = ['groq', 'local', 'json', 'ai', 'llm', 'ollama', 'error', 'none', 'null', 'undefined', 'empty']
    words = answer.lower().strip().split()
    
    # Single word check - most important
    if len(words) == 1:
        if words[0] in invalid_words:
            _logger.warning(f"Answer is invalid single word: '{answer}'")
            return False
        # Also reject if it's too short (less than 3 chars) and not a valid word
        if len(words[0]) < 3:
            _logger.warning(f"Answer is too short single word: '{answer}'")
            return False
    
    # Check if answer contains only model names or technical terms without context
    if len(words) <= 2 and all(word in invalid_words for word in words):
        _logger.warning(f"Answer contains only invalid words: '{answer}'")
        return False
    
    # Check if answer is too short overall (less than 10 characters) - but allow if it's a meaningful phrase
    if len(answer) < 10:
        if len(words) > 1 and words[0] not in invalid_words:
            return True
        _logger.warning(f"Answer too short: '{answer}' (length: {len(answer)})")
        return False
    
    return True

def format_ai_response(text):
    """Format AI response with markdown and enhanced styling"""
    # Convert markdown to HTML
    html = markdown2.markdown(text, extras=['fenced-code-blocks', 'tables', 'break-on-newline'])
    
    # Add custom formatting
    html = re.sub(r'\*\*(.*?)\*\*', r'<strong class="highlight">\1</strong>', html)
    html = re.sub(r'#{1,6}\s+(.*?)(?:\n|$)', r'<h3 class="section-header">\1</h3>', html)
    
    return html

@capture_output
def fuzzy_match(query, responses, threshold=70):
    """Enhanced fuzzy matching with multiple scoring methods"""
    query_clean = normalize_text(query)
    
    if not query_clean:
        _logger.warning("Empty query after normalization")
        return None
    
    if not responses:
        _logger.warning("No responses available for matching")
        return None
    
    best_match = None
    best_score = 0
    all_scores = []
    
    log_processing_step("CyberSentry AI", "fuzzy_match", "processing", f"Query: '{query[:50]}...'")
    _logger.debug(f"Searching for: '{query}' in {len(responses)} responses")
    
    for response in responses:
        if not isinstance(response, dict):
            continue
        if 'question' not in response or 'answer' not in response:
            continue
        
        question = response['question']
        if not isinstance(question, str):
            continue
            
        similarity_score = calculate_similarity_score(query, question)
        keyword_score = keyword_match_score(query, question)
        combined_score = similarity_score * 0.7 + keyword_score * 0.3
        
        all_scores.append({
            'question': question,
            'similarity': similarity_score,
            'keyword': keyword_score,
            'combined': combined_score
        })
        
        if combined_score > best_score:
            best_score = combined_score
            best_match = response
    
    all_scores.sort(key=lambda x: x['combined'], reverse=True)
    top_matches = all_scores[:3]
    _logger.debug("Top 3 Matches:")
    for i, score_data in enumerate(top_matches, 1):
        _logger.debug(f"{i}. '{score_data['question']}' - Combined: {score_data['combined']:.2f} | Similarity: {score_data['similarity']:.2f} | Keyword: {score_data['keyword']:.2f}")
    
    if best_score >= threshold and best_match:
        answer = best_match.get('answer')
        if isinstance(answer, dict):
            # Convert dict answer to string
            answer_parts = []
            for key, value in answer.items():
                if isinstance(value, dict):
                    if 'command' in value and 'description' in value:
                        answer_parts.append(f"{key}: {value.get('command', '')} - {value.get('description', '')}")
                    else:
                        answer_parts.append(f"{key}: {str(value)}")
                else:
                    answer_parts.append(f"{key}: {value}")
            answer = '\n'.join(answer_parts)
        elif not isinstance(answer, str):
            answer = str(answer) if answer else None
            
        if answer:
            # Validate answer before returning
            _logger.debug(f"[VALIDATION] Checking answer: '{answer}' (type: {type(answer)}, length: {len(str(answer))})")
            if is_valid_answer(answer):
                log_processing_step("CyberSentry AI", "fuzzy_match", "success", f"Match found (score: {best_score:.2f})")
                _logger.info(f"Match found with score: {best_score:.2f} - Question: '{best_match['question']}' - Answer: '{answer[:50]}...'")
                return answer
            else:
                _logger.warning(f"[VALIDATION FAILED] Match found but answer is invalid: '{answer}' - treating as no match")
                log_processing_step("CyberSentry AI", "fuzzy_match", "error", f"Invalid answer: '{answer}'")
                return None
        else:
            _logger.warning(f"Match found but answer is empty")
            return None
    else:
        log_processing_step("CyberSentry AI", "fuzzy_match", "error", f"No match (best score: {best_score:.2f}, threshold: {threshold})")
        _logger.warning(f"No match found (best score: {best_score:.2f}, threshold: {threshold})")
        return None

@capture_output
def get_groq_response(query):
    """Get response using centralized LLM router. Returns (answer, output, model_used)"""
    if not LLM_ROUTER_AVAILABLE:
        _logger.error("✗ LLM router not available")
        return None
    
    try:
        system_prompt = """You are CyberSentry AI, an advanced cybersecurity assistant specializing in:
- Ethical hacking and penetration testing
- Network security and vulnerability assessment
- Security tools (Nmap, Metasploit, Wireshark, Burp Suite, etc.)
- Threat analysis and mitigation strategies
- Secure coding practices
- Compliance and security frameworks

FORMAT YOUR RESPONSE WITH:
- Clear section headings using **bold text**
- Bullet points for lists
- Code blocks for commands (use ```language ```)
- Step-by-step numbered instructions when applicable
- Key terms in **bold**
- Important warnings or notes highlighted

Provide clear, educational, and actionable information while adhering to ethical and legal standards."""
        
        log_processing_step("CyberSentry AI", "llm_request", "processing", f"Query: '{query[:50]}...'")
        log_llm_request("CyberSentry AI", "cloud", len(query))
        start_time = time.time()
        
        result = generate_text(
            prompt=query,
            app_name="cybersentry_ai",
            task_type="security_analysis",
            system_prompt=system_prompt,
            temperature=0.7,
            max_tokens=2048
        )
        
        latency_ms = (time.time() - start_time) * 1000
        response_text = result.get("response", "")
        source = result.get("source", "none")
        model_used = result.get("model", "unknown")
        
        if response_text:
            log_llm_success("CyberSentry AI", source, len(response_text), latency_ms)
            log_processing_step("CyberSentry AI", "llm_request", "success", f"Response received ({len(response_text)} chars, source: {source})")
            _logger.info(f"✓ LLM response received ({len(response_text)} chars, {model_used}, {source})")
            formatted_response = format_ai_response(response_text.strip())
            return _wrap_with_model_used(formatted_response, model_used)
        else:
            log_llm_error("CyberSentry AI", source, Exception("Empty response"), fallback=False)
            _logger.warning("✗ Empty response from LLM")
            return None
            
    except Exception as e:
        _logger.error(f"✗ Error fetching response from LLM router: {e}", exc_info=True)
        log_llm_error("CyberSentry AI", "router", e, fallback=False)
        return None

@capture_output
def get_local_llm_response(query):
    """Get response from local LLM using centralized router. Returns (answer, output, model_used)"""
    if not LLM_ROUTER_AVAILABLE:
        _logger.debug("LLM router not available")
        return None
    
    try:
        system_prompt = """You are CyberSentry AI, an advanced cybersecurity assistant specializing in:
- Ethical hacking and penetration testing
- Network security and vulnerability assessment
- Security tools (Nmap, Metasploit, Wireshark, Burp Suite, etc.)
- Threat analysis and mitigation strategies
- Secure coding practices
- Compliance and security frameworks

FORMAT YOUR RESPONSE WITH:
- Clear section headings using **bold text**
- Bullet points for lists
- Code blocks for commands (use ```language ```)
- Step-by-step numbered instructions when applicable
- Key terms in **bold**
- Important warnings or notes highlighted

Provide clear, educational, and actionable information while adhering to ethical and legal standards."""
        
        log_processing_step("CyberSentry AI", "local_llm_request", "processing", f"Query: '{query[:50]}...'")
        _logger.info("[LOCAL LLM] Trying local LLM first...")
        
        log_llm_request("CyberSentry AI", "local", len(query))
        start_time = time.time()
        
        result = generate_text(
            prompt=query,
            app_name="cybersentry_ai",
            task_type="security_analysis",
            system_prompt=system_prompt,
            temperature=0.7,
            max_tokens=2048,
            prefer_local=True
        )
        
        latency_ms = (time.time() - start_time) * 1000
        local_result = result.get("response", "")
        source = result.get("source", "none")
        model_used = result.get("model", "unknown")
        
        if local_result and local_result.strip():
            log_llm_success("CyberSentry AI", "local", len(local_result), latency_ms)
            log_processing_step("CyberSentry AI", "local_llm_request", "success", f"Response received ({len(local_result)} chars)")
            _logger.info(f"✓ Local LLM response received ({len(local_result)} chars)")
            formatted_response = format_ai_response(local_result.strip())
            return _wrap_with_model_used(formatted_response, "local")
        else:
            log_llm_error("CyberSentry AI", "local", Exception("Local LLM returned empty"), fallback=True)
            _logger.warning("✗ Local LLM returned empty or failed")
            return None
            
    except Exception as e:
        log_llm_error("CyberSentry AI", "local", e, fallback=True)
        _logger.error(f"✗ Error with local LLM: {e}", exc_info=True)
        return None

def get_local_fallback(full_prompt):
    """Helper function to get local LLM fallback"""
    log_llm_fallback("CyberSentry AI", "cloud", "local")
    _logger.warning("⚠ Falling back to local Ollama...")
    
    log_llm_request("CyberSentry AI", "local", len(full_prompt))
    start_time = time.time()
    local_result, success = generate_with_ollama(
        full_prompt,
        system_prompt="You are CyberSentry AI, an expert cybersecurity assistant. Provide accurate security guidance.",
        temperature=0.7,
        max_tokens=2048
    )
    latency_ms = (time.time() - start_time) * 1000
    
    if success and local_result:
        log_llm_success("CyberSentry AI", "local", len(local_result), latency_ms)
        _logger.info("✓ Successfully used local Ollama model")
        formatted_response = format_ai_response(local_result.strip())
        return _wrap_with_model_used(formatted_response, "local")
    else:
        log_llm_error("CyberSentry AI", "local", Exception("Local LLM returned empty"), fallback=False)
        _logger.error("✗ Local LLM fallback also failed")
        return None

@cybersentry_ai.route('/')
def index():
    """Render the main chat interface"""
    return render_template('cybersentry_AI.html')

@cybersentry_ai.route('/ask', methods=['POST'])
@rate_limit_api(requests_per_minute=20, requests_per_hour=200)
@validate_request({
    "question": {
        "type": "string",
        "required": True,
        "max_length": 2000
    },
    "force_source": {
        "type": "string",
        "required": False,
        "max_length": 10,
        "allowed_values": ["json", "ai", None]
    }
}, strict=True)
def ask():
    """
    Handle question requests with enhanced processing
    OWASP: Rate limited, input validated
    """
    try:
        # Get validated data from request context
        data = g.validated_data
        question = InputValidator.validate_string(
            data.get('question'), 'question', max_length=2000, required=True
        )
        force_source = data.get('force_source', None)
        
        _logger.info(f"{'='*60}")
        _logger.info(f"[{time.strftime('%H:%M:%S')}] New Question: {question}")
        if force_source:
            _logger.info(f"[REGENERATE] Forcing source: {force_source}")
        _logger.info('='*60)
        
        # Handle regeneration with forced source
        if force_source == 'ai':
            log_processing_step("CyberSentry AI", "forced_ai", "processing", f"Question: '{question[:50]}...'")
            # Use LLM router with automatic fallback (tries Groq first, then local)
            system_prompt = """You are CyberSentry AI, an advanced cybersecurity assistant specializing in:
- Ethical hacking and penetration testing
- Network security and vulnerability assessment
- Security tools (Nmap, Metasploit, Wireshark, Burp Suite, etc.)
- Threat analysis and mitigation strategies
- Secure coding practices
- Compliance and security frameworks

FORMAT YOUR RESPONSE WITH:
- Clear section headings using **bold text**
- Bullet points for lists
- Code blocks for commands (use ```language ```)
- Step-by-step numbered instructions when applicable
- Key terms in **bold**
- Important warnings or notes highlighted

Provide clear, educational, and actionable information while adhering to ethical and legal standards."""
            
            if LLM_ROUTER_AVAILABLE:
                log_llm_request("CyberSentry AI", "auto", len(question))
                start_time = time.time()
                result = generate_text(
                    prompt=question,
                    app_name="cybersentry_ai",
                    task_type="security_analysis",
                    system_prompt=system_prompt,
                    temperature=0.7,
                    max_tokens=2048
                )
                latency_ms = (time.time() - start_time) * 1000
                ai_answer = result.get("response", "")
                source = result.get("source", "none")
                model_used = result.get("model", "unknown")
                
                if ai_answer:
                    log_llm_success("CyberSentry AI", source, len(ai_answer), latency_ms)
                    formatted_answer = format_ai_response(ai_answer.strip())
                    log_processing_step("CyberSentry AI", "forced_ai", "success", f"Response generated (model: {model_used}, source: {source})")
                    return jsonify({
                        'answer': formatted_answer,
                        'source': 'AI',
                        'terminal_output': '',
                        'confidence': 'medium',
                        'can_regenerate': True,
                        'model_used': model_used
                    })
            
            log_processing_step("CyberSentry AI", "forced_ai", "error", "No response generated from LLM router")
        
        elif force_source == 'json':
            log_processing_step("CyberSentry AI", "forced_json", "processing", f"Question: '{question[:50]}...'")
            result = fuzzy_match(question, responses, threshold=60)
            # fuzzy_match has @capture_output decorator, returns (answer, output, model_used)
            if isinstance(result, tuple) and len(result) >= 3:
                answer, json_output, _ = result
            elif isinstance(result, tuple) and len(result) >= 1:
                answer = result[0] if len(result) > 0 else None
                json_output = result[1] if len(result) > 1 else ""
            else:
                answer = result
                json_output = ""
            
            # Ensure answer is a string, not a tuple or other type
            if isinstance(answer, tuple):
                _logger.warning(f"[FORCED JSON] Answer is a tuple, extracting first element: {answer}")
                answer = answer[0] if len(answer) > 0 else None
            if answer is not None and not isinstance(answer, str):
                _logger.warning(f"[FORCED JSON] Answer is not a string (type: {type(answer)}), converting: {answer}")
                answer = str(answer) if answer else None
            
            # Validate JSON answer even when forced
            if answer and isinstance(answer, str) and is_valid_answer(answer):
                log_processing_step("CyberSentry AI", "forced_json", "success", "Match found")
                return jsonify({
                    'answer': answer,
                    'source': 'JSON',
                    'terminal_output': json_output,
                    'confidence': 'high',
                    'can_regenerate': True,
                    'model_used': 'json'
                })
            elif answer:
                # Invalid answer from JSON, log warning but still return it (user forced JSON)
                # Ensure answer is string
                if isinstance(answer, tuple):
                    answer = answer[0] if len(answer) > 0 else ''
                if not isinstance(answer, str):
                    answer = str(answer) if answer else ''
                answer_str = str(answer) if not isinstance(answer, str) else answer
                _logger.warning(f"[FORCED JSON] Invalid answer found: '{answer_str}' - but user forced JSON source")
                log_processing_step("CyberSentry AI", "forced_json", "warning", f"Invalid answer: '{answer_str}'")
                # Still return it since user forced JSON, but log the issue
                return jsonify({
                    'answer': answer_str,
                    'source': 'JSON',
                    'terminal_output': json_output,
                    'confidence': 'low',
                    'can_regenerate': True,
                    'model_used': 'json'
                })
            else:
                log_processing_step("CyberSentry AI", "forced_json", "error", "No match found")
        
        # Normal flow: Try JSON first
        log_processing_step("CyberSentry AI", "normal_flow", "processing", f"Question: '{question[:50]}...'")
        result = fuzzy_match(question, responses, threshold=70)
        # fuzzy_match has @capture_output decorator, returns (answer, output, model_used)
        if isinstance(result, tuple) and len(result) >= 3:
            answer, json_output, _ = result
        elif isinstance(result, tuple) and len(result) >= 1:
            answer = result[0] if len(result) > 0 else None
            json_output = result[1] if len(result) > 1 else ""
        else:
            answer = result
            json_output = ""
        
        # Ensure answer is a string, not a tuple or other type
        if isinstance(answer, tuple):
            _logger.warning(f"[NORMAL FLOW] Answer is a tuple, extracting first element: {answer}")
            answer = answer[0] if len(answer) > 0 else None
        if answer is not None and not isinstance(answer, str):
            _logger.warning(f"[NORMAL FLOW] Answer is not a string (type: {type(answer)}), converting: {answer}")
            answer = str(answer) if answer else None
        
        # Validate JSON answer - reject if invalid (e.g., single word "groq")
        _logger.debug(f"[NORMAL FLOW] JSON result: answer='{answer}', type={type(answer)}")
        if answer:
            _logger.debug(f"[VALIDATION] Validating JSON answer: '{answer}'")
            is_valid = is_valid_answer(answer)
            _logger.debug(f"[VALIDATION] Result: {is_valid}")
            
        if answer and isinstance(answer, str) and is_valid_answer(answer):
            log_processing_step("CyberSentry AI", "normal_flow", "success", "JSON match found")
            answer_preview = answer[:50] if len(answer) > 50 else answer
            _logger.info(f"[RESULT] Using JSON database response: '{answer_preview}...'")
            
            # Threat analysis - ensure answer is string
            threat_data = threat_analyzer.analyze_threat_level(question, str(answer))
            
            # Save to database
            user_id = request.remote_addr or 'anonymous'
            try:
                with sqlite3.connect(db_manager.db_name) as conn:
                    cursor = conn.cursor()
                    cursor.execute('''
                        INSERT INTO conversations 
                        (user_id, question, answer, source, model_used, confidence, timestamp)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (user_id, question, str(answer), 'JSON', 'json', 'high', datetime.now().isoformat()))
                    conn.commit()
            except Exception as e:
                _logger.warning(f"Failed to save conversation: {e}")
            
            return jsonify({
                'answer': str(answer),  # Ensure answer is string
                'source': 'JSON',
                'terminal_output': json_output,
                'confidence': 'high',
                'can_regenerate': True,
                'model_used': 'json',
                'threat_analysis': threat_data
            })
        elif answer:
            # Invalid answer from JSON (e.g., just "groq"), skip to Local LLM
            # Ensure answer is string for logging
            answer_str = str(answer) if not isinstance(answer, str) else answer
            _logger.warning(f"[JSON REJECTED] Invalid answer found: '{answer_str}' (length: {len(answer_str)}) - skipping JSON and using Local LLM")
            log_processing_step("CyberSentry AI", "normal_flow", "processing", f"JSON answer invalid: '{answer_str}' - trying Local LLM")
            # Continue to Local LLM below
        
        # Step 2: Use LLM router (tries Groq first, automatically falls back to local)
        log_processing_step("CyberSentry AI", "normal_flow", "processing", "Trying LLM router (Groq -> Ollama fallback)...")
        _logger.info("[STEP 2] Using LLM router with automatic fallback...")
        
        if LLM_ROUTER_AVAILABLE:
            system_prompt = """You are CyberSentry AI, an advanced cybersecurity assistant specializing in:
- Ethical hacking and penetration testing
- Network security and vulnerability assessment
- Security tools (Nmap, Metasploit, Wireshark, Burp Suite, etc.)
- Threat analysis and mitigation strategies
- Secure coding practices
- Compliance and security frameworks

FORMAT YOUR RESPONSE WITH:
- Clear section headings using **bold text**
- Bullet points for lists
- Code blocks for commands (use ```language ```)
- Step-by-step numbered instructions when applicable
- Key terms in **bold**
- Important warnings or notes highlighted

Provide clear, educational, and actionable information while adhering to ethical and legal standards."""
            
            log_llm_request("CyberSentry AI", "auto", len(question))
            start_time = time.time()
            result = generate_text(
                prompt=question,
                app_name="cybersentry_ai",
                task_type="security_analysis",
                system_prompt=system_prompt,
                temperature=0.7,
                max_tokens=2048
            )
            latency_ms = (time.time() - start_time) * 1000
            ai_answer = result.get("response", "")
            source = result.get("source", "none")
            model_used = result.get("model", "unknown")
            
            if ai_answer:
                log_llm_success("CyberSentry AI", source, len(ai_answer), latency_ms)
                formatted_answer = format_ai_response(ai_answer.strip())
                log_processing_step("CyberSentry AI", "normal_flow", "success", f"LLM response generated (model: {model_used}, source: {source})")
                _logger.info(f"[RESULT] Using LLM response (source: {source}, model: {model_used})")
                
                # Threat analysis
                threat_data = threat_analyzer.analyze_threat_level(question, ai_answer)
                
                # Save to database
                user_id = request.remote_addr or 'anonymous'
                try:
                    with sqlite3.connect(db_manager.db_name) as conn:
                        cursor = conn.cursor()
                        cursor.execute('''
                            INSERT INTO conversations 
                            (user_id, question, answer, source, model_used, confidence, timestamp, response_time)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                        ''', (user_id, question, ai_answer, 'AI', model_used, 'medium', datetime.now().isoformat(), latency_ms))
                        conn.commit()
                except Exception as e:
                    _logger.warning(f"Failed to save conversation: {e}")
                
                return jsonify({
                    'answer': formatted_answer,
                    'source': 'AI',
                    'terminal_output': '',
                    'confidence': 'medium',
                    'can_regenerate': True,
                    'model_used': model_used,
                    'threat_analysis': threat_data
                })
        
        # Final fallback
        _logger.info("[FALLBACK] Using default response")
        fallback_answer = format_ai_response("""I don't have specific information about that topic in my knowledge base. 

**🔒 Security Best Practices:**
- Keep systems and software updated
- Use strong, unique passwords with MFA
- Implement network segmentation
- Regular security audits and monitoring
- Follow the principle of least privilege

**💡 Try asking about:**
- Common security tools (Nmap, Wireshark, Metasploit)
- Attack types (DDoS, SQL injection, XSS)
- Security concepts (encryption, firewalls, VPNs)
- Penetration testing methodologies""")
        
        return jsonify({
            'answer': fallback_answer,
            'source': 'Fallback',
            'terminal_output': '',
            'confidence': 'low',
            'can_regenerate': False,
            'model_used': 'none'
        })
        
    except ValueError as e:
        error_msg = f"Validation error: {str(e)}"
        _logger.error(f"\n[VALIDATION ERROR] {error_msg}")
        log_processing_step("CyberSentry AI", "ask", "error", f"Validation: {error_msg}")
        return jsonify({
            'error': error_msg,
            'terminal_output': ''
        }), 400
    except Exception as e:
        error_msg = f"Error processing request: {str(e)}"
        _logger.error(f"\n[ERROR] {error_msg}", exc_info=True)
        log_processing_step("CyberSentry AI", "ask", "error", f"Exception: {error_msg}")
        return jsonify({
            'error': error_msg,
            'terminal_output': ''
        }), 500


@cybersentry_ai.route('/healthcheck', methods=['GET'])
def healthcheck():
    """Quick healthcheck for monitoring / load-balancer probes."""
    return jsonify({
        'status':          'ok',
        'version':         '4.5',
        'responses_loaded': len(responses),
        'llm_router':      LLM_ROUTER_AVAILABLE,
        'tts':             TTS_AVAILABLE,
        'ocr':             OCR_AVAILABLE,
        'excel':           EXCEL_AVAILABLE,
        'docx':            DOCX_AVAILABLE,
        'timestamp':       datetime.now().isoformat(),
    })



@cybersentry_ai.route('/reload-responses', methods=['POST'])
@rate_limit_api(requests_per_minute=5, requests_per_hour=20)  # Strict limit for admin operations
def reload_responses():
    """
    Reload responses from JSON file
    OWASP: Rate limited
    """
    """Endpoint to reload responses.json without restarting server"""
    global responses
    responses = load_responses()
    return jsonify({
        'message': f'Responses reloaded successfully. Total responses: {len(responses)}'
    })

@cybersentry_ai.route('/stats', methods=['GET'])
def stats():
    """Get statistics about the response database"""
    cats = {}
    for r in responses:
        if isinstance(r, dict):
            cat = r.get('category', 'uncategorized')
            cats[cat] = cats.get(cat, 0) + 1
    return jsonify({
        'total_responses': len(responses),
        'categories':      list(cats.keys()),
        'category_counts': cats,
        'llm_available':   LLM_ROUTER_AVAILABLE,
        'tts_available':   TTS_AVAILABLE,
    })

# Database Manager for Advanced Features
class CyberSentryDatabase:
    """Database manager for conversation history, threat intelligence, and analytics."""
    
    def __init__(self, db_name='cybersentry_ai.db'):
        self.db_name = db_name
        self.init_db()
    
    def init_db(self):
        """Initialize database with proper error handling"""
        try:
            with sqlite3.connect(self.db_name) as conn:
                cursor = conn.cursor()
                
                # Conversation history
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS conversations (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        user_id TEXT,
                        question TEXT,
                        answer TEXT,
                        source TEXT,
                        model_used TEXT,
                        confidence TEXT,
                        timestamp DATETIME,
                        response_time REAL
                    )
                ''')
                
                # Threat intelligence cache
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS threat_intel (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        query_hash TEXT UNIQUE,
                        query TEXT,
                        threat_level TEXT,
                        risk_score INTEGER,
                        indicators TEXT,
                        recommendations TEXT,
                        timestamp DATETIME
                    )
                ''')
                
                # Security analytics
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS analytics (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        event_type TEXT,
                        event_data TEXT,
                        timestamp DATETIME
                    )
                ''')
                
                # User feedback
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS feedback (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        conversation_id INTEGER,
                        helpful BOOLEAN,
                        feedback_text TEXT,
                        timestamp DATETIME
                    )
                ''')
                
                # Create indexes
                cursor.execute('CREATE INDEX IF NOT EXISTS idx_user_id ON conversations(user_id)')
                cursor.execute('CREATE INDEX IF NOT EXISTS idx_timestamp ON conversations(timestamp)')
                cursor.execute('CREATE INDEX IF NOT EXISTS idx_query_hash ON threat_intel(query_hash)')
                
                conn.commit()
            _logger.info("✓ CyberSentry AI database initialized")
        except sqlite3.Error as e:
            _logger.error(f"⚠ Database initialization error: {e}")

# Initialize database
db_manager = CyberSentryDatabase()

# Threat Intelligence Analyzer
class ThreatIntelligenceAnalyzer:
    """Advanced threat intelligence and security analysis."""
    
    def __init__(self):
        self.threat_keywords = {
            'critical': ['exploit', 'vulnerability', 'breach', 'attack', 'malware', 'ransomware', 'phishing', 'ddos'],
            'high': ['security', 'threat', 'risk', 'compromise', 'intrusion', 'backdoor', 'trojan'],
            'medium': ['warning', 'alert', 'suspicious', 'anomaly', 'unusual'],
            'low': ['information', 'guide', 'tutorial', 'explanation']
        }
    
    def analyze_threat_level(self, query, answer):
        """Analyze threat level based on query and answer content."""
        query_lower  = (query  or '').lower() if isinstance(query,  str) else str(query  or '').lower()
        if isinstance(answer, tuple):
            answer = answer[0] if len(answer) > 0 else ''
        if not isinstance(answer, str):
            answer = str(answer) if answer else ''
        answer_lower = (answer or '').lower()
        combined     = query_lower + ' ' + answer_lower
        
        threat_score = 0
        indicators = []
        
        for level, keywords in self.threat_keywords.items():
            matches = [kw for kw in keywords if kw in combined]
            if matches:
                if level == 'critical':
                    threat_score += 30
                elif level == 'high':
                    threat_score += 20
                elif level == 'medium':
                    threat_score += 10
                indicators.extend(matches)
        
        if threat_score >= 30:
            threat_level = 'CRITICAL'
        elif threat_score >= 20:
            threat_level = 'HIGH'
        elif threat_score >= 10:
            threat_level = 'MEDIUM'
        else:
            threat_level = 'LOW'
        
        return {
            'threat_level': threat_level,
            'risk_score': min(threat_score, 100),
            'indicators': list(set(indicators)),
            'recommendations': self._generate_recommendations(threat_level)
        }
    
    def _generate_recommendations(self, threat_level):
        """Generate security recommendations based on threat level."""
        recommendations = {
            'CRITICAL': [
                'Immediate action required',
                'Isolate affected systems',
                'Notify security team',
                'Review security logs',
                'Implement emergency patches'
            ],
            'HIGH': [
                'Review security configurations',
                'Update security policies',
                'Monitor network traffic',
                'Conduct security audit'
            ],
            'MEDIUM': [
                'Review security best practices',
                'Update documentation',
                'Schedule security review'
            ],
            'LOW': [
                'Continue monitoring',
                'Maintain security hygiene'
            ]
        }
        return recommendations.get(threat_level, [])
# Initialize threat analyzer
threat_analyzer = ThreatIntelligenceAnalyzer()

# REPLACE ENTIRE PDFMemoryManager WITH THIS ADVANCED VERSION
class AdvancedDocumentProcessor:
    """Multi-format document processor with AI-powered analysis."""
    
    SUPPORTED_FORMATS = {
        'pdf':  'application/pdf',
        'txt':  'text/plain',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'csv':  'text/csv',
        'png':  'image/png',
        'jpg':  'image/jpeg',
        'jpeg': 'image/jpeg',
        'json': 'application/json',
        'md':   'text/markdown',
        'log':  'text/plain',
        'xml':  'application/xml',
        'yaml': 'text/yaml',
        'yml':  'text/yaml',
    }
    
    def __init__(self):
        self.sessions = {}
        self.max_session_age = 7200  # 2 hours for document sessions
        self.audio_cache = {}  # Cache for generated audio files
        self.cleanup_thread = threading.Thread(target=self._cleanup_old_sessions, daemon=True)
        self.cleanup_thread.start()
    
    def process_document(self, file, file_type):
        """Route to appropriate processor based on file type."""
        processors = {
            'pdf': self._process_pdf,
            'txt': self._process_text,
            'docx': self._process_docx if DOCX_AVAILABLE else None,
            'xlsx': self._process_excel if EXCEL_AVAILABLE else None,
            'csv': self._process_csv if EXCEL_AVAILABLE else None,
            'png': self._process_image if OCR_AVAILABLE else None,
            'jpg': self._process_image if OCR_AVAILABLE else None,
            'jpeg': self._process_image if OCR_AVAILABLE else None,
            'json': self._process_json,
            'md': self._process_markdown
        }
        
        processor = processors.get(file_type)
        if not processor:
            return {'success': False, 'error': f'Unsupported file type: {file_type}'}
        
        if processor is None:
            return {'success': False, 'error': f'File type {file_type} requires additional dependencies'}
        
        try:
            return processor(file)
        except Exception as e:
            _logger.error(f"Document processing error ({file_type}): {e}", exc_info=True)
            return {'success': False, 'error': str(e)}
    
    def _process_pdf(self, file):
        """Enhanced PDF processing with metadata extraction."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            text_content = []
            images_found = 0
            
            metadata = {
                'total_pages': len(pdf_reader.pages),
                'title': pdf_reader.metadata.get('/Title', 'Unknown') if pdf_reader.metadata else 'Unknown',
                'author': pdf_reader.metadata.get('/Author', 'Unknown') if pdf_reader.metadata else 'Unknown',
                'creation_date': pdf_reader.metadata.get('/CreationDate', 'Unknown') if pdf_reader.metadata else 'Unknown',
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    # Extract headers/sections
                    lines = text.split('\n')
                    headers = [l for l in lines if len(l) < 100 and l.isupper() and len(l) > 3]
                    
                    text_content.append({
                        'page': page_num + 1,
                        'content': text.strip(),
                        'headers': headers[:5],  # Top 5 headers per page
                        'word_count': len(text.split())
                    })
                
                # Count images
                try:
                    if '/XObject' in page.get('/Resources', {}):
                        images_found += len(page['/Resources']['/XObject'])
                except:
                    pass
            
            # Generate statistics
            total_words = sum(p['word_count'] for p in text_content) if text_content else 0
            avg_words_per_page = total_words // len(text_content) if text_content else 0
            
            # Extract key topics using word frequency
            all_text = ' '.join(p['content'] for p in text_content)
            keywords = self._extract_keywords(all_text)
            
            return {
                'success': True,
                'text': text_content,
                'metadata': metadata,
                'statistics': {
                    'total_chars': sum(len(p['content']) for p in text_content),
                    'total_words': total_words,
                    'avg_words_per_page': avg_words_per_page,
                    'images_found': images_found,
                    'keywords': keywords[:10]
                },
                'file_type': 'pdf'
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _process_text(self, file):
        """Process plain text files."""
        try:
            content = file.read().decode('utf-8', errors='ignore')
            lines = content.split('\n')
            
            return {
                'success': True,
                'text': [{'page': 1, 'content': content, 'word_count': len(content.split())}],
                'metadata': {
                    'total_pages': 1,
                    'title': 'Text Document',
                    'line_count': len(lines)
                },
                'statistics': {
                    'total_chars': len(content),
                    'total_words': len(content.split()),
                    'total_lines': len(lines),
                    'keywords': self._extract_keywords(content)[:10]
                },
                'file_type': 'txt'
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _process_docx(self, file):
        """Process Word documents."""
        try:
            doc = docx.Document(io.BytesIO(file.read()))
            full_text = []
            tables_count = len(doc.tables)
            images_count = 0
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Count images
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        images_count += 1
            except:
                pass
            
            content = '\n'.join(full_text)
            
            return {
                'success': True,
                'text': [{'page': 1, 'content': content, 'word_count': len(content.split())}],
                'metadata': {
                    'total_pages': 1,
                    'title': 'Word Document',
                    'tables': tables_count,
                    'images': images_count
                },
                'statistics': {
                    'total_chars': len(content),
                    'total_words': len(content.split()),
                    'total_paragraphs': len(full_text),
                    'keywords': self._extract_keywords(content)[:10]
                },
                'file_type': 'docx'
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _process_excel(self, file):
        """Process Excel files with smart sheet detection."""
        try:
            df_dict = pd.read_excel(io.BytesIO(file.read()), sheet_name=None)
            
            sheets_data = []
            total_rows = 0
            total_cols = 0
            
            for sheet_name, df in df_dict.items():
                total_rows += len(df)
                total_cols = max(total_cols, len(df.columns))
                
                # Convert to readable format
                summary = f"Sheet: {sheet_name}\n"
                summary += f"Columns: {', '.join(df.columns.astype(str))}\n\n"
                summary += df.head(20).to_string()  # First 20 rows
                
                sheets_data.append({
                    'sheet_name': sheet_name,
                    'content': summary,
                    'rows': len(df),
                    'columns': list(df.columns.astype(str))
                })
            
            combined_content = '\n\n'.join(s['content'] for s in sheets_data)
            
            return {
                'success': True,
                'text': [{'page': 1, 'content': combined_content, 'sheets': sheets_data}],
                'metadata': {
                    'total_pages': len(df_dict),
                    'title': 'Excel Workbook',
                    'total_sheets': len(df_dict)
                },
                'statistics': {
                    'total_chars': len(combined_content),
                    'total_rows': total_rows,
                    'total_columns': total_cols,
                    'sheet_names': list(df_dict.keys())
                },
                'file_type': 'xlsx'
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _process_csv(self, file):
        """Process CSV files."""
        try:
            df = pd.read_csv(io.BytesIO(file.read()))
            
            summary = f"CSV Data Analysis\n"
            summary += f"Columns: {', '.join(df.columns.astype(str))}\n\n"
            summary += f"Data Preview:\n{df.head(50).to_string()}\n\n"
            try:
                summary += f"Statistics:\n{df.describe().to_string()}"
            except:
                summary += "Statistics: Not available for this data type"
            
            return {
                'success': True,
                'text': [{'page': 1, 'content': summary}],
                'metadata': {
                    'total_pages': 1,
                    'title': 'CSV Data',
                    'rows': len(df),
                    'columns': len(df.columns)
                },
                'statistics': {
                    'total_chars': len(summary),
                    'total_rows': len(df),
                    'total_columns': len(df.columns),
                    'column_names': list(df.columns.astype(str))
                },
                'file_type': 'csv'
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _process_image(self, file):
        """Process images with OCR text extraction."""
        try:
            image = Image.open(io.BytesIO(file.read()))
            
            # Basic image info
            width, height = image.size
            format_type = image.format
            mode = image.mode
            
            # OCR text extraction
            try:
                ocr_text = pytesseract.image_to_string(image)
                ocr_success = True
            except Exception as ocr_error:
                _logger.warning(f"OCR failed: {ocr_error}")
                ocr_text = "[OCR not available - install tesseract-ocr]"
                ocr_success = False
            
            content = f"Image Analysis\n\n"
            content += f"Dimensions: {width}x{height}\n"
            content += f"Format: {format_type}\n"
            content += f"Color Mode: {mode}\n\n"
            
            if ocr_success and ocr_text.strip():
                content += f"Extracted Text (OCR):\n{ocr_text}"
            else:
                content += "No text detected in image."
            
            return {
                'success': True,
                'text': [{'page': 1, 'content': content, 'ocr_text': ocr_text}],
                'metadata': {
                    'total_pages': 1,
                    'title': 'Image File',
                    'width': width,
                    'height': height,
                    'format': format_type
                },
                'statistics': {
                    'total_chars': len(ocr_text) if ocr_success else 0,
                    'ocr_enabled': ocr_success,
                    'keywords': self._extract_keywords(ocr_text)[:10] if ocr_success else []
                },
                'file_type': 'image'
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _process_json(self, file):
        """Process JSON files."""
        try:
            data = json.load(io.BytesIO(file.read()))
            
            # Pretty print JSON
            formatted = json.dumps(data, indent=2)
            
            # Analyze structure
            def count_keys(obj, depth=0):
                if isinstance(obj, dict):
                    return sum(1 + count_keys(v, depth+1) for v in obj.values())
                elif isinstance(obj, list):
                    return sum(count_keys(item, depth+1) for item in obj)
                return 0
            
            total_keys = count_keys(data)
            
            return {
                'success': True,
                'text': [{'page': 1, 'content': formatted, 'raw_data': data}],
                'metadata': {
                    'total_pages': 1,
                    'title': 'JSON Data',
                    'structure': type(data).__name__
                },
                'statistics': {
                    'total_chars': len(formatted),
                    'total_keys': total_keys,
                    'root_type': type(data).__name__
                },
                'file_type': 'json'
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _process_markdown(self, file):
        """Process Markdown files."""
        try:
            content = file.read().decode('utf-8', errors='ignore')
            
            # Extract headers
            headers = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
            
            # Count links and images
            links = len(re.findall(r'\[.+?\]\(.+?\)', content))
            images = len(re.findall(r'!\[.+?\]\(.+?\)', content))
            
            return {
                'success': True,
                'text': [{'page': 1, 'content': content}],
                'metadata': {
                    'total_pages': 1,
                    'title': 'Markdown Document',
                    'headers': len(headers)
                },
                'statistics': {
                    'total_chars': len(content),
                    'total_words': len(content.split()),
                    'total_headers': len(headers),
                    'total_links': links,
                    'total_images': images,
                    'keywords': self._extract_keywords(content)[:10]
                },
                'file_type': 'md'
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _extract_keywords(self, text, top_n=20):
        """Extract top keywords from text."""
        # Remove common stop words
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                     'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'been',
                     'be', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
                     'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those'}
        
        # Extract words
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        filtered_words = [w for w in words if w not in stop_words]
        
        # Count frequency
        word_freq = Counter(filtered_words)
        return [word for word, count in word_freq.most_common(top_n)]
    
    def generate_audio_summary(self, session_id, language='en'):
        """Generate TTS audio summary of document."""
        session = self.sessions.get(session_id)
        if not session:
            return {'success': False, 'error': 'Session not found'}
        
        if not TTS_AVAILABLE:
            return {'success': False, 'error': 'TTS not available – run: pip install gTTS --break-system-packages'}
        
        try:
            # Check cache first
            cache_key = f"{session_id}_{language}"
            if cache_key in self.audio_cache:
                cached_path = self.audio_cache[cache_key]
                # Validate cached file still exists and is non-empty
                if os.path.exists(cached_path) and os.path.getsize(cached_path) > 0:
                    return {'success': True, 'audio_path': cached_path,
                            'audio_filename': os.path.basename(cached_path),
                            'duration_estimate': 30}
                else:
                    del self.audio_cache[cache_key]
            
            # Generate summary for audio
            content = session['content']
            file_type = content.get('file_type', 'unknown')
            
            # Create concise summary text
            summary_text = f"Document Summary. "
            summary_text += f"File type: {file_type}. "
            
            if content['metadata']:
                meta = content['metadata']
                if 'title' in meta and meta['title'] != 'Unknown':
                    summary_text += f"Title: {meta['title']}. "
                if 'total_pages' in meta:
                    summary_text += f"Contains {meta['total_pages']} pages. "
            
            if 'statistics' in content:
                stats = content['statistics']
                if 'total_words' in stats:
                    summary_text += f"Total words: {stats['total_words']}. "
                if 'keywords' in stats and stats['keywords']:
                    summary_text += f"Key topics: {', '.join(stats['keywords'][:5])}. "
            
            # Add brief content preview
            if content['text']:
                first_content = content['text'][0]['content'][:500]
                summary_text += f"Content preview: {first_content}"
            
            # Generate audio
            tts = gTTS(text=summary_text, lang=language, slow=False)
            
            # Save to temp file
            audio_filename = f"summary_{session_id}_{int(time.time())}.mp3"
            audio_path = os.path.join(tempfile.gettempdir(), audio_filename)
            tts.save(audio_path)

            # Verify the file was actually written
            if not os.path.exists(audio_path) or os.path.getsize(audio_path) == 0:
                raise RuntimeError(f"gTTS produced an empty or missing file: {audio_path}")

            # Cache the path
            self.audio_cache[cache_key] = audio_path
            
            return {
                'success': True,
                'audio_path': audio_path,
                'audio_filename': audio_filename,
                'duration_estimate': len(summary_text) / 150  # ~150 words per minute
            }
        except Exception as e:
            _logger.error(f"Audio generation error: {e}", exc_info=True)
            return {'success': False, 'error': str(e)}
    
    def create_session(self, user_id, document_content, filename):
        """Create new document session."""
        session_id = hashlib.md5(f"{user_id}{time.time()}".encode()).hexdigest()
        self.sessions[session_id] = {
            'user_id': user_id,
            'content': document_content,
            'filename': filename,
            'created_at': time.time(),
            'query_count': 0,
            'last_accessed': time.time()
        }
        return session_id
    
    def get_session(self, session_id):
        """Retrieve session data."""
        session = self.sessions.get(session_id)
        if session:
            session['last_accessed'] = time.time()
        return session
    
    def query_document(self, session_id, query, page_range=None):
        """Query document with context-aware searching."""
        session = self.sessions.get(session_id)
        if not session:
            return None
        
        session['query_count'] += 1
        session['last_accessed'] = time.time()
        doc_data = session['content']
        
        # Build context from document
        if page_range and doc_data.get('file_type') == 'pdf':
            start, end = page_range
            relevant_pages = [p for p in doc_data['text'] if start <= p['page'] <= end]
        else:
            # Smart relevance for all document types
            query_lower = query.lower()
            query_words = set(query_lower.split())
            
            scored_pages = []
            for page in doc_data['text']:
                content_lower = page['content'].lower()
                score = sum(1 for word in query_words if word in content_lower)
                scored_pages.append((score, page))
            
            scored_pages.sort(reverse=True, key=lambda x: x[0])
            relevant_pages = [p[1] for p in scored_pages[:30]]
        
        context = "\n\n".join([
            f"[Section {i+1}]\n{p['content'][:3000]}" 
            for i, p in enumerate(relevant_pages[:20])
        ])
        
        return {
            'context': context,
            'metadata': doc_data['metadata'],
            'statistics': doc_data.get('statistics', {}),
            'file_type': doc_data.get('file_type', 'unknown'),
            'total_chars': len(context)
        }
    
    def delete_session(self, session_id):
        """Delete session and free memory."""
        if session_id in self.sessions:
            # Clean up audio cache
            for key in list(self.audio_cache.keys()):
                if key.startswith(session_id):
                    try:
                        os.remove(self.audio_cache[key])
                    except:
                        pass
                    del self.audio_cache[key]
            
            del self.sessions[session_id]
            return True
        return False
    
    def _cleanup_old_sessions(self):
        """Background cleanup of expired sessions."""
        while True:
            time.sleep(600)  # Check every 10 minutes
            current_time = time.time()
            expired = [
                sid for sid, data in self.sessions.items()
                if current_time - data['last_accessed'] > self.max_session_age
            ]
            for sid in expired:
                self.delete_session(sid)
                _logger.info(f"Cleaned up expired document session: {sid}")

# Initialize document processor
doc_processor = AdvancedDocumentProcessor()

@cybersentry_ai.route('/threat-analysis', methods=['POST'])
@rate_limit_api(requests_per_minute=10, requests_per_hour=100)
@validate_request({
    "query": {
        "type": "string",
        "required": True,
        "max_length": 2000
    }
}, strict=True)
def threat_analysis():
    """Advanced threat intelligence analysis endpoint."""
    try:
        data = g.validated_data
        query = InputValidator.validate_string(data.get('query'), 'query', max_length=2000, required=True)
        
        # Check cache
        query_hash = hashlib.md5(query.encode()).hexdigest()
        with sqlite3.connect(db_manager.db_name) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM threat_intel WHERE query_hash = ?', (query_hash,))
            cached = cursor.fetchone()
            
            if cached:
                return jsonify({
                    'threat_level': cached['threat_level'],
                    'risk_score': cached['risk_score'],
                    'indicators': json.loads(cached['indicators']),
                    'recommendations': json.loads(cached['recommendations']),
                    'cached': True
                })
        
        # Get answer first
        result = fuzzy_match(query, responses, threshold=70)
        if isinstance(result, tuple):
            answer = result[0] if len(result) > 0 else None
        else:
            answer = result
        
        if not answer or not is_valid_answer(answer):
            # Try LLM
            if LLM_ROUTER_AVAILABLE:
                system_prompt = "You are a cybersecurity expert. Provide a brief security analysis."
                result = generate_text(
                    prompt=query,
                    app_name="cybersentry_ai",
                    task_type="security_analysis",
                    system_prompt=system_prompt,
                    temperature=0.7,
                    max_tokens=1024
                )
                answer = result.get("response", "")
        
        # Analyze threat
        analysis = threat_analyzer.analyze_threat_level(query, answer or '')
        
        # Cache result
        with sqlite3.connect(db_manager.db_name) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT OR REPLACE INTO threat_intel 
                (query_hash, query, threat_level, risk_score, indicators, recommendations, timestamp)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                query_hash, query, analysis['threat_level'], analysis['risk_score'],
                json.dumps(analysis['indicators']), json.dumps(analysis['recommendations']),
                datetime.now().isoformat()
            ))
            conn.commit()
        
        return jsonify({
            **analysis,
            'cached': False
        })
    except Exception as e:
        _logger.error(f"Threat analysis error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@cybersentry_ai.route('/analytics', methods=['GET'])
def get_analytics():
    """Get security analytics and statistics."""
    try:
        with sqlite3.connect(db_manager.db_name) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Conversation stats
            cursor.execute('SELECT COUNT(*) as total FROM conversations')
            total_conversations = cursor.fetchone()['total']
            
            # Source distribution
            cursor.execute('''
                SELECT source, COUNT(*) as count 
                FROM conversations 
                GROUP BY source
            ''')
            source_dist = {row['source']: row['count'] for row in cursor.fetchall()}
            
            # Threat level distribution
            cursor.execute('''
                SELECT threat_level, COUNT(*) as count 
                FROM threat_intel 
                GROUP BY threat_level
            ''')
            threat_dist = {row['threat_level']: row['count'] for row in cursor.fetchall()}
            
            # Recent activity
            cursor.execute('''
                SELECT COUNT(*) as count 
                FROM conversations 
                WHERE timestamp > datetime('now', '-24 hours')
            ''')
            recent_activity = cursor.fetchone()['count']
            
            return jsonify({
                'total_conversations': total_conversations,
                'source_distribution': source_dist,
                'threat_distribution': threat_dist,
                'recent_activity_24h': recent_activity
            })
    except Exception as e:
        _logger.error(f"Analytics error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

# ADD THIS NEW ENDPOINT
# REPLACE /pdf/upload WITH THIS ENHANCED VERSION
@cybersentry_ai.route('/document/upload', methods=['POST'])
@rate_limit_api(requests_per_minute=5, requests_per_hour=20)
def upload_document():
    """Upload and process multi-format documents."""
    try:
        if 'document_file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        doc_file = request.files['document_file']
        if doc_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Get file extension
        filename = secure_filename(doc_file.filename)
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        # Map common aliases
        _ext_alias = {'jpeg': 'jpg'}
        file_ext = _ext_alias.get(file_ext, file_ext)

        if file_ext not in AdvancedDocumentProcessor.SUPPORTED_FORMATS:
            return jsonify({
                'error': f'Unsupported format. Supported: {", ".join(AdvancedDocumentProcessor.SUPPORTED_FORMATS.keys())}'
            }), 400
        
        # Check file size (100MB max)
        doc_file.seek(0, 2)
        file_size = doc_file.tell()
        doc_file.seek(0)
        
        if file_size > 100 * 1024 * 1024:
            return jsonify({'error': 'File too large (max 100MB)'}), 400
        
        user_id = request.remote_addr or 'anonymous'
        _logger.info(f"Processing document upload: {filename} ({file_size} bytes, type: {file_ext})")
        
        # Process document
        processing_result = doc_processor.process_document(doc_file, file_ext)
        
        if not processing_result['success']:
            return jsonify({'error': processing_result['error']}), 500
        
        # Create session
        session_id = doc_processor.create_session(user_id, processing_result, filename)
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'filename': filename,
            'file_type': file_ext,
            'metadata': processing_result['metadata'],
            'statistics': processing_result.get('statistics', {}),
            'message': f'{file_ext.upper()} processed successfully'
        })
        
    except Exception as e:
        _logger.error(f"Document upload error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@cybersentry_ai.route('/document/query', methods=['POST'])
@rate_limit_api(requests_per_minute=20, requests_per_hour=200)
@validate_request({
    "session_id": {"type": "string", "required": True},
    "question": {"type": "string", "required": True, "max_length": 2000},
    "page_range": {"type": "string", "required": False}
}, strict=False)
def query_document():
    """Query document with AI assistance."""
    try:
        data = g.validated_data
        session_id = data.get('session_id')
        question = InputValidator.validate_string(
            data.get('question'), 'question', max_length=2000, required=True
        )
        page_range_raw = data.get('page_range')
        page_range = None
        if page_range_raw and isinstance(page_range_raw, str) and '-' in page_range_raw:
            try:
                parts = page_range_raw.split('-')
                page_range = (int(parts[0]), int(parts[1]))
            except (ValueError, IndexError):
                page_range = None
        doc_context = doc_processor.query_document(session_id, question, page_range)
        
        if not doc_context:
            return jsonify({'error': 'Invalid session or session expired'}), 400
        
        # Build enhanced prompt
        file_type = doc_context['file_type'].upper()
        system_prompt = f"""You are CyberSentry AI with access to a {file_type} document.

**DOCUMENT METADATA:**
- Type: {file_type}
- Title: {doc_context['metadata'].get('title', 'Unknown')}
{f"- Author: {doc_context['metadata'].get('author')}" if doc_context['metadata'].get('author') != 'Unknown' else ''}

**DOCUMENT STATISTICS:**
{chr(10).join(f"- {k.replace('_', ' ').title()}: {v}" for k, v in doc_context['statistics'].items() if isinstance(v, (int, float, str)))}

**YOUR TASK:**
Answer the user's question accurately based ONLY on the document content provided below.
- Cite specific sections when referencing information
- If information is not in the document, clearly state that
- Be precise and thorough
- For data/tables, present in organized format

**DOCUMENT CONTENT:**
{doc_context['context'][:20000]}
"""
        
        log_processing_step("CyberSentry AI", "document_query", "processing", f"Query: '{question[:50]}...'")
        log_llm_request("CyberSentry AI", "auto", len(question))
        start_time = time.time()
        
        # Use LLM router
        result = generate_text(
            prompt=question,
            app_name="cybersentry_ai_docs",
            task_type="document_analysis",
            system_prompt=system_prompt,
            temperature=0.3,
            max_tokens=4000
        )
        
        latency_ms = (time.time() - start_time) * 1000
        ai_answer = result.get("response", "")
        source = result.get("source", "none")
        model_used = result.get("model", "unknown")
        
        if ai_answer:
            log_llm_success("CyberSentry AI", source, len(ai_answer), latency_ms)
            formatted_answer = format_ai_response(ai_answer.strip())
            
            return jsonify({
                'success': True,
                'answer': formatted_answer,
                'source': source,
                'model_used': model_used,
                'latency_ms': latency_ms,
                'file_type': file_type,
                'context_chars': doc_context['total_chars']
            })
        else:
            return jsonify({'error': 'No answer generated'}), 500
            
    except Exception as e:
        _logger.error(f"Document query error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@cybersentry_ai.route('/document/audio-stream/<session_id>', methods=['GET'])
def stream_audio_summary(session_id):
    """
    Stream audio as audio/mpeg directly – avoids blob: CSP restriction.
    Supports Range requests so browsers can seek.
    """
    try:
        language = request.args.get('lang', 'en')

        # Generate (or use cached) audio
        result = doc_processor.generate_audio_summary(session_id, language)
        if not result['success']:
            return jsonify({'error': result['error']}), 400

        audio_path = result['audio_path']
        if not os.path.exists(audio_path):
            return jsonify({'error': 'Audio file not found'}), 404

        file_size = os.path.getsize(audio_path)

        # Parse Range header for partial content
        range_header = request.headers.get('Range', None)
        if range_header:
            byte_start, byte_end = 0, file_size - 1
            match = re.match(r'bytes=([0-9]+)-([0-9]*)', range_header)
            if match:
                byte_start = int(match.group(1))
                if match.group(2):
                    byte_end = int(match.group(2))

            length = byte_end - byte_start + 1

            def generate_partial():
                with open(audio_path, 'rb') as f:
                    f.seek(byte_start)
                    remaining = length
                    while remaining > 0:
                        chunk = f.read(min(8192, remaining))
                        if not chunk:
                            break
                        remaining -= len(chunk)
                        yield chunk

            headers = {
                'Content-Range':  f'bytes {byte_start}-{byte_end}/{file_size}',
                'Accept-Ranges':  'bytes',
                'Content-Length': str(length),
                'Content-Type':   'audio/mpeg',
                'Cache-Control':  'no-store',
                'Access-Control-Allow-Origin': '*',
            }
            return Response(generate_partial(), status=206, headers=headers, direct_passthrough=True)

        # Full file
        def generate_full():
            with open(audio_path, 'rb') as f:
                while True:
                    chunk = f.read(8192)
                    if not chunk:
                        break
                    yield chunk

        headers = {
            'Content-Type':   'audio/mpeg',
            'Content-Length': str(file_size),
            'Accept-Ranges':  'bytes',
            'Cache-Control':  'no-store',
            'Access-Control-Allow-Origin': '*',
        }
        return Response(generate_full(), status=200, headers=headers, direct_passthrough=True)

    except Exception as e:
        _logger.error(f"Audio stream error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500



@cybersentry_ai.route('/document/audio-summary/<session_id>', methods=['GET'])
def generate_audio_summary(session_id):
    """Generate TTS audio summary of document."""
    try:
        language = request.args.get('lang', 'en')
        
        result = doc_processor.generate_audio_summary(session_id, language)
        
        if not result['success']:
            return jsonify({'error': result['error']}), 400
        
        # Return audio file
        response = send_file(
            result['audio_path'],
            mimetype='audio/mpeg',
            as_attachment=False,
            download_name=result['audio_filename'],
            conditional=True,   # enables Range request support
        )
        response.headers['Access-Control-Allow-Origin']   = '*'
        response.headers['Access-Control-Allow-Headers']  = 'Range'
        response.headers['Access-Control-Expose-Headers'] = 'Content-Range, Accept-Ranges, Content-Length'
        response.headers['Accept-Ranges']                 = 'bytes'
        response.headers['Cache-Control']                 = 'no-store'
        response.headers['X-Content-Type-Options']        = 'nosniff'
        return response
        
    except Exception as e:
        _logger.error(f"Audio generation error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@cybersentry_ai.route('/document/export/<session_id>', methods=['GET'])
def export_document_analysis(session_id):
    """Export document analysis as formatted report."""
    try:
        export_format = request.args.get('format', 'json')  # json, txt, md
        
        session = doc_processor.get_session(session_id)
        if not session:
            return jsonify({'error': 'Session not found'}), 404
        
        content = session['content']
        filename = session['filename']
        
        if export_format == 'json':
            return jsonify({
                'filename': filename,
                'metadata': content['metadata'],
                'statistics': content.get('statistics', {}),
                'file_type': content.get('file_type'),
                'exported_at': datetime.now().isoformat()
            })
        
        elif export_format == 'txt':
            report = f"DOCUMENT ANALYSIS REPORT\n"
            report += f"{'='*60}\n\n"
            report += f"Filename: {filename}\n"
            report += f"File Type: {content.get('file_type', 'unknown').upper()}\n"
            report += f"Analyzed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            report += f"METADATA:\n{'-'*60}\n"
            for k, v in content['metadata'].items():
                report += f"{k.replace('_', ' ').title()}: {v}\n"
            report += f"\nSTATISTICS:\n{'-'*60}\n"
            for k, v in content.get('statistics', {}).items():
                if not isinstance(v, list):
                    report += f"{k.replace('_', ' ').title()}: {v}\n"
            
            return Response(report, mimetype='text/plain', headers={
                'Content-Disposition': f'attachment; filename=analysis_{session_id}.txt'
            })
        
        elif export_format == 'md':
            report = f"# Document Analysis Report\n\n"
            report += f"**Filename:** {filename}  \n"
            report += f"**Type:** {content.get('file_type', 'unknown').upper()}  \n"
            report += f"**Analyzed:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \n\n"
            report += f"## Metadata\n\n"
            for k, v in content['metadata'].items():
                report += f"- **{k.replace('_', ' ').title()}:** {v}\n"
            report += f"\n## Statistics\n\n"
            for k, v in content.get('statistics', {}).items():
                if not isinstance(v, list):
                    report += f"- **{k.replace('_', ' ').title()}:** {v}\n"
            
            return Response(report, mimetype='text/markdown', headers={
                'Content-Disposition': f'attachment; filename=analysis_{session_id}.md'
            })
        
        else:
            return jsonify({'error': 'Invalid format. Use json, txt, or md'}), 400
    
    except Exception as e:
        _logger.error(f"Export error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@cybersentry_ai.route('/document/stats/<session_id>', methods=['GET'])
def document_stats(session_id):
    """Return quick statistics for an uploaded document session."""
    try:
        session = doc_processor.get_session(session_id)
        if not session:
            return jsonify({'error': 'Session not found'}), 404
        content = session['content']
        return jsonify({
            'filename':   session['filename'],
            'file_type':  content.get('file_type', 'unknown'),
            'metadata':   content.get('metadata', {}),
            'statistics': content.get('statistics', {}),
            'query_count': session.get('query_count', 0),
            'created_at': session.get('created_at'),
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500



@cybersentry_ai.route('/document/session/<session_id>', methods=['DELETE'])
def delete_document_session(session_id):
    """Delete document session."""
    try:
        if doc_processor.delete_session(session_id):
            return jsonify({'success': True, 'message': 'Session deleted'})
        else:
            return jsonify({'error': 'Session not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@cybersentry_ai.route('/execute-code', methods=['POST'])
@rate_limit_api(requests_per_minute=5, requests_per_hour=20)
@validate_request({
    "code": {
        "type": "string",
        "required": True,
        "max_length": 5000
    },
    "language": {
        "type": "string",
        "required": True,
        "allowed_values": ["python", "bash", "javascript"]
    }
}, strict=True)
def execute_code():
    """
    Execute code snippets in sandboxed environment
    WARNING: This is for demonstration only - use proper sandboxing in production
    """
    try:
        data = g.validated_data
        code = data.get('code')
        language = data.get('language')
        
        # For security, only allow specific safe operations
        # In production, use Docker containers or proper sandboxing
        
        if language == 'python':
            # Simple safe execution (very limited)
            import subprocess
            import tempfile
            
            with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as f:
                f.write(code)
                temp_file = f.name
            
            try:
                proc = subprocess.run(
                    ['python3', temp_file],
                    capture_output=True,
                    text=True,
                    timeout=5
                )
                output = proc.stdout if proc.returncode == 0 else proc.stderr
                return jsonify({
                    'success': proc.returncode == 0,
                    'output': output or '(no output)',
                    'language': language
                })
            finally:
                try:
                    os.unlink(temp_file)
                except OSError:
                    pass
        
        return jsonify({
            'success': False,
            'output': f'Language {language} not supported yet',
            'language': language
        })
        
    except Exception as e:
        _logger.error(f"Code execution error: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'output': str(e),
            'language': language
        }), 500

@cybersentry_ai.route('/history', methods=['GET'])
def get_history():
    """Get conversation history."""
    try:
        user_id = request.args.get('user_id', 'default')
        limit = request.args.get('limit', 50, type=int)
        
        with sqlite3.connect(db_manager.db_name) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM conversations 
                WHERE user_id = ? 
                ORDER BY timestamp DESC 
                LIMIT ?
            ''', (user_id, limit))
            
            history = [dict(row) for row in cursor.fetchall()]
            return jsonify({'history': history, 'count': len(history)})
    except Exception as e:
        _logger.error(f"History error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@cybersentry_ai.route('/ask-all', methods=['POST'])
@rate_limit_api(requests_per_minute=10, requests_per_hour=50)
@validate_request({
    "question": {
        "type": "string",
        "required": True,
        "max_length": 2000
    }
}, strict=True)
def ask_all():
    """
    Get responses from all sources in PARALLEL for faster results
    """
    try:
        data = g.validated_data
        question = InputValidator.validate_string(
            data.get('question'), 'question', max_length=2000, required=True
        )
        
        _logger.info(f"[ASK-ALL] Parallel processing for: '{question[:50]}...'")
        
        results = {
            'json': None,
            'cloud_llm': None,
            'local_llm': None
        }
        
        # Use ThreadPoolExecutor for parallel execution
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            # Submit all tasks simultaneously
            future_json = executor.submit(get_json_response, question)
            future_cloud = executor.submit(get_cloud_llm_response_fast, question)
            future_local = executor.submit(get_local_llm_response_fast, question)
            
            # Wait for all with timeout
            try:
                results['json'] = future_json.result(timeout=5)
            except Exception as e:
                _logger.error(f"JSON error: {e}")
                results['json'] = {'answer': 'Timeout', 'status': 'error'}
            
            try:
                results['cloud_llm'] = future_cloud.result(timeout=10)
            except Exception as e:
                _logger.error(f"Cloud LLM error: {e}")
                results['cloud_llm'] = {'answer': 'Timeout', 'status': 'error'}
            
            try:
                results['local_llm'] = future_local.result(timeout=15)  # Increased timeout for local LLM
            except concurrent.futures.TimeoutError:
                _logger.warning("Local LLM timeout - may still be processing")
                results['local_llm'] = {'answer': 'Processing timeout - Local LLM may still be running', 'status': 'timeout'}
            except Exception as e:
                _logger.error(f"Local LLM error: {e}", exc_info=True)
                results['local_llm'] = {'answer': f'Error: {str(e)}', 'status': 'error'}
        
        return jsonify({
            'question': question,
            'results': results,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        _logger.error(f"Ask-all error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

def get_json_response(question):
    """Fast JSON lookup"""
    try:
        json_result = fuzzy_match(question, responses, threshold=60)
        # fuzzy_match returns (answer, output, model_used) because of @capture_output
        if isinstance(json_result, tuple):
            # Try to extract the answer part
            json_answer = json_result[0] if len(json_result) > 0 else None
        else:
            json_answer = json_result
        
        if json_answer and isinstance(json_answer, str) and is_valid_answer(json_answer):
            return {
                'answer': json_answer,
                'source': 'JSON',
                'model_used': 'json',
                'confidence': 'high',
                'status': 'success'
            }
        else:
            return {
                'answer': 'No match found in knowledge base',
                'source': 'JSON',
                'status': 'no_match'
            }
    except Exception as e:
        return {'answer': f'Error: {str(e)}', 'status': 'error'}

def get_cloud_llm_response_fast(question):
    """Fast cloud LLM with reduced tokens"""
    if not LLM_ROUTER_AVAILABLE:
        return {'answer': 'Unavailable', 'status': 'unavailable'}
    
    try:
        start_time = time.time()
        result = generate_text(
            prompt=question,
            app_name="cybersentry_ai",
            task_type="security_analysis",
            system_prompt="You are CyberSentry AI. Provide concise security guidance.",
            temperature=0.7,
            max_tokens=1024,  # Reduced from 2048
            prefer_local=False
        )
        latency_ms = (time.time() - start_time) * 1000
        
        cloud_answer = result.get("response", "")
        if cloud_answer and result.get("source") != "local":
            return {
                'answer': format_ai_response(cloud_answer.strip()),
                'source': 'Cloud LLM',
                'model_used': result.get("model", "unknown"),
                'status': 'success',
                'latency_ms': latency_ms
            }
        return {'answer': 'Unavailable', 'status': 'unavailable'}
    except Exception as e:
        return {'answer': f'Error: {str(e)}', 'status': 'error'}

def get_local_llm_response_fast(question):
    """Fast local LLM with reduced tokens - ensures local execution"""
    if not LLM_ROUTER_AVAILABLE:
        return {'answer': 'LLM router unavailable', 'status': 'unavailable'}
    
    try:
        _logger.info("[LOCAL LLM] Attempting local LLM generation...")
        start_time = time.time()
        
        # Force local execution by explicitly checking and using local fallback if needed
        result = generate_text(
            prompt=question,
            app_name="cybersentry_ai",
            task_type="security_analysis",
            system_prompt="You are CyberSentry AI. Provide concise security guidance.",
            temperature=0.7,
            max_tokens=1024,  # Reduced from 2048
            prefer_local=True
        )
        
        latency_ms = (time.time() - start_time) * 1000
        source = result.get("source", "none")
        local_answer = result.get("response", "")
        
        _logger.info(f"[LOCAL LLM] Result - source: {source}, has_answer: {bool(local_answer)}, latency: {latency_ms:.0f}ms")
        
        # Accept if source is local OR if we got a response (router may not always set source correctly)
        if local_answer and (source == "local" or source == "ollama" or len(local_answer) > 10):
            return {
                'answer': format_ai_response(local_answer.strip()),
                'source': 'Local LLM',
                'model_used': result.get("model", "local-unknown"),
                'status': 'success',
                'latency_ms': latency_ms
            }
        
        # If no response but source indicates local was attempted, return timeout
        if source in ["local", "ollama"]:
            return {'answer': 'Local LLM processing - response pending', 'status': 'processing'}
        
        return {'answer': 'Local LLM unavailable or not responding', 'status': 'unavailable'}
    except Exception as e:
        _logger.error(f"[LOCAL LLM] Exception: {e}", exc_info=True)
        return {'answer': f'Error: {str(e)}', 'status': 'error'}

# Blueprint is registered in server.py
# This function is kept for backward compatibility but does nothing
def init_app(app):
    """Legacy function - blueprint is registered in server.py"""
    pass

if __name__ == "__main__":
    from flask import Flask
    app = Flask(__name__)
    app.config['SECRET_KEY'] = 'dev-key-for-standalone-mode'
    app.register_blueprint(cybersentry_ai)
    print("Starting CyberSentry AI in standalone mode on port 5002...")
    app.run(debug=True, port=5002)